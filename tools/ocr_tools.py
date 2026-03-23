"""OCR tools: extract text from images and PDFs using EasyOCR."""
import io
import json
from pathlib import Path
from typing import Dict, List, Optional

from mcp.server.fastmcp import FastMCP

from utils.file_utils import resolve_input, resolve_output, make_output_path
from utils.ocr_utils import get_ocr_reader, OCR_LANGUAGES, bbox_to_points
from utils.pdf_utils import open_pdf, render_page, pixmap_to_pil


def register_ocr_tools(app: FastMCP) -> None:

    # ------------------------------------------------------------------ #
    #  ocr_image                                                            #
    # ------------------------------------------------------------------ #
    @app.tool()
    def ocr_image(
        input_path: str,
        languages: Optional[List[str]] = None,
        output_format: str = "text",
        output_path: Optional[str] = None,
        gpu: bool = False,
    ) -> Dict:
        """
        Extract text from an image using OCR (EasyOCR).

        languages: list of language codes, e.g. ['en'] or ['en', 'fr'].
                   Use list_ocr_languages to see all supported codes.
                   Default: ['en'].
        output_format: 'text' (plain string) or 'json' (with bounding boxes + confidence).
        output_path: optional file path to save results (.txt or .json).
        gpu: use GPU acceleration if available (default False).

        Note: First call per language set downloads model weights (~100-400 MB).
        """
        try:
            src = resolve_input(input_path)
            langs = languages or ["en"]

            reader = get_ocr_reader(langs, gpu=gpu)
            results = reader.readtext(str(src))
            # results: list of (bbox, text, confidence)

            if output_format == "json":
                data = [
                    {
                        "text": text,
                        "confidence": round(conf, 4),
                        "bbox": bbox,
                    }
                    for bbox, text, conf in results
                ]
                output = json.dumps(data, ensure_ascii=False, indent=2)
                ext = ".json"
            else:
                output = "\n".join(text for _, text, _ in results)
                ext = ".txt"

            if output_path:
                dst = resolve_output(output_path)
                dst.write_text(output, encoding="utf-8")
                saved_to = str(dst)
            elif len(output) > 2000:
                # Auto-save if result is large
                saved_to = make_output_path(str(src), "_ocr", ext)
                Path(saved_to).write_text(output, encoding="utf-8")
            else:
                saved_to = None

            return {
                "success": True,
                "text": output if output_format == "text" else None,
                "results": data if output_format == "json" else None,
                "word_count": len([r for r in results]),
                "languages": langs,
                "saved_to": saved_to,
            }
        except FileNotFoundError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"OCR failed: {e}"}

    # ------------------------------------------------------------------ #
    #  ocr_pdf                                                              #
    # ------------------------------------------------------------------ #
    @app.tool()
    def ocr_pdf(
        input_path: str,
        output_path: Optional[str] = None,
        languages: Optional[List[str]] = None,
        dpi: int = 200,
        create_searchable_pdf: bool = True,
        gpu: bool = False,
        pages: Optional[List[int]] = None,
    ) -> Dict:
        """
        Run OCR on a PDF (useful for scanned/image-based PDFs).

        create_searchable_pdf: if True, produces a new PDF with an invisible text
                               layer so the text becomes searchable/copyable.
                               If False, returns extracted text only.
        dpi: rendering resolution for OCR (higher = more accurate, slower).
             Recommended: 200-300.
        pages: 1-based list of pages to process (default: all).

        Note: First call per language set downloads EasyOCR model weights.
        """
        try:
            from PIL import Image
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas as rl_canvas
            import fitz

            src = resolve_input(input_path)
            langs = languages or ["en"]

            if not output_path:
                suffix = "_searchable" if create_searchable_pdf else "_ocr_text"
                ext = ".pdf" if create_searchable_pdf else ".txt"
                output_path = make_output_path(str(src), suffix, ext)
            dst = resolve_output(output_path)

            reader = get_ocr_reader(langs, gpu=gpu)
            doc = open_pdf(str(src))
            n = doc.page_count
            targets = [p - 1 for p in pages] if pages else list(range(n))
            targets = [i for i in targets if 0 <= i < n]

            all_text_pages = []

            if create_searchable_pdf:
                # Build an overlay PDF with invisible text, then merge
                overlay_buf = io.BytesIO()
                c = rl_canvas.Canvas(overlay_buf)

                for page_idx in targets:
                    pix = render_page(doc, page_idx, dpi=dpi)
                    img = pixmap_to_pil(pix)
                    page_w_pt = pix.width * 72.0 / dpi
                    page_h_pt = pix.height * 72.0 / dpi

                    c.setPageSize((page_w_pt, page_h_pt))

                    # Draw original page as background image
                    img_buf = io.BytesIO()
                    if img.mode == "RGBA":
                        img = img.convert("RGB")
                    img.save(img_buf, format="JPEG", quality=90)
                    img_buf.seek(0)
                    c.drawImage(
                        img_buf, 0, 0, width=page_w_pt, height=page_h_pt,
                        preserveAspectRatio=False
                    )

                    # OCR and add invisible text layer
                    results = reader.readtext(img_buf.getvalue() if False else str(src))
                    # Re-run OCR on the PIL image bytes
                    img_buf2 = io.BytesIO()
                    img.save(img_buf2, format="PNG")
                    results = reader.readtext(img_buf2.getvalue())

                    page_texts = []
                    c.setFillColorRGB(1, 1, 1, 0)  # fully transparent
                    c.setFont("Helvetica", 8)

                    for bbox, text, conf in results:
                        x0_pt, y0_pt, x1_pt, y1_pt = bbox_to_points(bbox, dpi)
                        # PDF Y-axis: flip (PDF origin is bottom-left)
                        pdf_y = page_h_pt - y1_pt
                        c.drawString(x0_pt, pdf_y, text)
                        page_texts.append(text)

                    all_text_pages.append("\n".join(page_texts))
                    c.showPage()

                c.save()

                # Merge overlay onto original PDF
                overlay_buf.seek(0)
                overlay_doc = fitz.open("pdf", overlay_buf.read())
                out_doc = fitz.open()

                for i, page_idx in enumerate(targets):
                    orig_page = doc[page_idx]
                    new_page = out_doc.new_page(
                        width=orig_page.rect.width,
                        height=orig_page.rect.height
                    )
                    new_page.show_pdf_page(new_page.rect, doc, page_idx)
                    if i < len(overlay_doc):
                        new_page.show_pdf_page(new_page.rect, overlay_doc, i)

                out_doc.save(str(dst), garbage=3, deflate=True)
                out_doc.close()
                overlay_doc.close()

            else:
                # Text-only output
                for page_idx in targets:
                    pix = render_page(doc, page_idx, dpi=dpi)
                    img = pixmap_to_pil(pix)
                    img_buf = io.BytesIO()
                    if img.mode == "RGBA":
                        img = img.convert("RGB")
                    img.save(img_buf, format="PNG")
                    results = reader.readtext(img_buf.getvalue())
                    all_text_pages.append("\n".join(text for _, text, _ in results))

                full_text = "\n\n--- Page Break ---\n\n".join(all_text_pages)
                dst.write_text(full_text, encoding="utf-8")

            doc.close()

            total_text = "\n".join(all_text_pages)
            return {
                "success": True,
                "output_path": str(dst),
                "pages_processed": len(targets),
                "languages": langs,
                "searchable_pdf": create_searchable_pdf,
                "extracted_text_preview": total_text[:500] + ("..." if len(total_text) > 500 else ""),
                "char_count": len(total_text),
            }
        except FileNotFoundError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"OCR PDF failed: {e}"}

    # ------------------------------------------------------------------ #
    #  image_to_word                                                        #
    # ------------------------------------------------------------------ #
    @app.tool()
    def image_to_word(
        input_path: str,
        output_path: Optional[str] = None,
        languages: Optional[List[str]] = None,
        gpu: bool = False,
    ) -> Dict:
        """
        Convert an image (JPG, PNG, etc.) to a Word (.docx) document using OCR.

        Extracts text from the image via EasyOCR and saves it as a formatted .docx file.
        languages: list of language codes, e.g. ['en']. Default: ['en'].

        Note: First call per language set downloads EasyOCR model weights (~100-400 MB).
        """
        try:
            from docx import Document

            src = resolve_input(input_path)
            langs = languages or ["en"]
            if not output_path:
                output_path = make_output_path(str(src), "_ocr", ".docx")
            dst = resolve_output(output_path)

            reader = get_ocr_reader(langs, gpu=gpu)
            results = reader.readtext(str(src))

            document = Document()
            document.add_heading(src.stem, level=1)
            for _, text, _ in results:
                if text.strip():
                    document.add_paragraph(text)
            document.save(str(dst))

            extracted_text = "\n".join(t for _, t, _ in results)
            return {
                "success": True,
                "output_path": str(dst),
                "char_count": len(extracted_text),
                "languages": langs,
                "file_size_bytes": dst.stat().st_size,
            }
        except FileNotFoundError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"image_to_word failed: {e}"}

    # ------------------------------------------------------------------ #
    #  image_to_html / image_to_markdown                                    #
    # ------------------------------------------------------------------ #
    @app.tool()
    def image_to_html(
        input_path: str,
        output_path: Optional[str] = None,
        languages: Optional[List[str]] = None,
        gpu: bool = False,
    ) -> Dict:
        """
        Extract text from an image via OCR and save as an HTML file.
        Bounding-box positions are used to approximate layout.
        """
        import html as html_module
        try:
            src = resolve_input(input_path)
            langs = languages or ["en"]
            if not output_path:
                output_path = make_output_path(str(src), "_ocr", ".html")
            dst = resolve_output(output_path)

            reader = get_ocr_reader(langs, gpu=gpu)
            results = reader.readtext(str(src))

            lines = [f'<!DOCTYPE html>\n<html lang="en"><head><meta charset="UTF-8">'
                     f'<title>{html_module.escape(src.stem)}</title>'
                     f'<style>body{{font-family:sans-serif;max-width:900px;margin:auto;padding:2em}}'
                     f'p{{margin:0.3em 0}}</style></head><body>']
            for _, text, _ in results:
                lines.append(f"<p>{html_module.escape(text)}</p>")
            lines.append("</body></html>")

            dst.write_text("\n".join(lines), encoding="utf-8")
            return {"success": True, "output_path": str(dst), "text_blocks": len(results), "languages": langs}
        except FileNotFoundError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"image_to_html failed: {e}"}

    @app.tool()
    def image_to_markdown(
        input_path: str,
        output_path: Optional[str] = None,
        languages: Optional[List[str]] = None,
        gpu: bool = False,
    ) -> Dict:
        """
        Extract text from an image via OCR and save as a Markdown file.
        Each OCR result block becomes a paragraph.
        """
        try:
            src = resolve_input(input_path)
            langs = languages or ["en"]
            if not output_path:
                output_path = make_output_path(str(src), "_ocr", ".md")
            dst = resolve_output(output_path)

            reader = get_ocr_reader(langs, gpu=gpu)
            results = reader.readtext(str(src))

            lines = [f"# {src.stem}\n"]
            lines.extend(text for _, text, _ in results if text.strip())
            content = "\n\n".join(lines)
            dst.write_text(content, encoding="utf-8")
            return {"success": True, "output_path": str(dst), "text_blocks": len(results), "languages": langs}
        except FileNotFoundError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"image_to_markdown failed: {e}"}

    # ------------------------------------------------------------------ #
    #  list_ocr_languages                                                   #
    # ------------------------------------------------------------------ #
    @app.tool()
    def list_ocr_languages() -> Dict:
        """
        List language codes supported by the OCR engine (EasyOCR).

        Pass one or more of these codes in the 'languages' parameter
        of ocr_pdf or ocr_image.
        """
        return {
            "success": True,
            "languages": OCR_LANGUAGES,
            "usage_example": ["en"],
            "note": (
                "First use of a new language set downloads model weights (~100-400 MB) "
                "to ~/.EasyOCR/. Pre-download with: "
                "python -c \"import easyocr; easyocr.Reader(['en'])\""
            ),
        }
