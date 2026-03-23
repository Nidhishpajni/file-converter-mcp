"""E-book and RTF conversion tools: EPUB ↔ TXT/HTML/MD/DOCX/PDF, RTF → TXT/HTML."""
import html as html_module
from pathlib import Path
from typing import Dict, List, Optional

from mcp.server.fastmcp import FastMCP
from utils.file_utils import resolve_input, resolve_output, make_output_path


def register_ebook_tools(app: FastMCP) -> None:

    def _epub_to_text(src: Path) -> str:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(str(src))
        parts = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "lxml")
            text = soup.get_text(separator="\n")
            if text.strip():
                parts.append(text.strip())
        return "\n\n".join(parts)

    def _epub_to_html_str(src: Path) -> str:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(str(src))
        title = book.get_metadata("DC", "title")
        title_str = title[0][0] if title else src.stem
        sections = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            soup = BeautifulSoup(item.get_content(), "lxml")
            body = soup.find("body")
            sections.append(str(body) if body else str(soup))
        return (
            f'<!DOCTYPE html>\n<html lang="en"><head><meta charset="UTF-8">'
            f'<title>{html_module.escape(title_str)}</title>'
            f'<style>body{{font-family:serif;max-width:800px;margin:auto;padding:2em}}</style>'
            f'</head><body>' + "\n".join(sections) + "</body></html>"
        )

    # ── EPUB → TXT ─────────────────────────────────────────────────────────
    @app.tool()
    def epub_to_txt(input_path: str, output_path: Optional[str] = None) -> Dict:
        """Extract plain text from an EPUB e-book."""
        try:
            src = resolve_input(input_path)
            if not output_path:
                output_path = make_output_path(input_path, "", ".txt")
            dst = resolve_output(output_path)
            text = _epub_to_text(src)
            dst.write_text(text, encoding="utf-8")
            return {"success": True, "output_path": str(dst), "char_count": len(text)}
        except ImportError:
            return {"success": False, "error": "ebooklib not installed. Run: pip install ebooklib"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ── EPUB → HTML ────────────────────────────────────────────────────────
    @app.tool()
    def epub_to_html(input_path: str, output_path: Optional[str] = None) -> Dict:
        """Convert an EPUB e-book to a single HTML file."""
        try:
            src = resolve_input(input_path)
            if not output_path:
                output_path = make_output_path(input_path, "", ".html")
            dst = resolve_output(output_path)
            html_content = _epub_to_html_str(src)
            dst.write_text(html_content, encoding="utf-8")
            return {"success": True, "output_path": str(dst), "file_size_bytes": dst.stat().st_size}
        except ImportError:
            return {"success": False, "error": "ebooklib not installed. Run: pip install ebooklib"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ── EPUB → Markdown ────────────────────────────────────────────────────
    @app.tool()
    def epub_to_markdown(input_path: str, output_path: Optional[str] = None) -> Dict:
        """Convert an EPUB e-book to Markdown."""
        try:
            from bs4 import BeautifulSoup
            import ebooklib
            from ebooklib import epub

            src = resolve_input(input_path)
            if not output_path:
                output_path = make_output_path(input_path, "", ".md")
            dst = resolve_output(output_path)

            book = epub.read_epub(str(src))
            lines = []
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                soup = BeautifulSoup(item.get_content(), "lxml")
                for tag in soup.find_all(["h1","h2","h3","p","li"]):
                    text = tag.get_text(strip=True)
                    if not text:
                        continue
                    if tag.name == "h1": lines.append(f"# {text}")
                    elif tag.name == "h2": lines.append(f"## {text}")
                    elif tag.name == "h3": lines.append(f"### {text}")
                    elif tag.name == "li": lines.append(f"- {text}")
                    else: lines.append(text)
                lines.append("")

            content = "\n".join(lines)
            dst.write_text(content, encoding="utf-8")
            return {"success": True, "output_path": str(dst), "char_count": len(content)}
        except ImportError:
            return {"success": False, "error": "ebooklib not installed. Run: pip install ebooklib"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ── EPUB → DOCX ────────────────────────────────────────────────────────
    @app.tool()
    def epub_to_docx(input_path: str, output_path: Optional[str] = None) -> Dict:
        """Convert an EPUB e-book to a Word document (.docx)."""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            from docx import Document

            src = resolve_input(input_path)
            if not output_path:
                output_path = make_output_path(input_path, "", ".docx")
            dst = resolve_output(output_path)

            book = epub.read_epub(str(src))
            doc = Document()
            title_meta = book.get_metadata("DC", "title")
            if title_meta:
                doc.add_heading(title_meta[0][0], level=1)

            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                soup = BeautifulSoup(item.get_content(), "lxml")
                for tag in soup.find_all(["h1","h2","h3","p","li"]):
                    text = tag.get_text(strip=True)
                    if not text:
                        continue
                    if tag.name == "h1": doc.add_heading(text, level=1)
                    elif tag.name == "h2": doc.add_heading(text, level=2)
                    elif tag.name == "h3": doc.add_heading(text, level=3)
                    elif tag.name == "li": doc.add_paragraph(text, style="List Bullet")
                    else: doc.add_paragraph(text)

            doc.save(str(dst))
            return {"success": True, "output_path": str(dst), "file_size_bytes": dst.stat().st_size}
        except ImportError:
            return {"success": False, "error": "ebooklib not installed. Run: pip install ebooklib"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ── EPUB → PDF ─────────────────────────────────────────────────────────
    @app.tool()
    def epub_to_pdf(input_path: str, output_path: Optional[str] = None) -> Dict:
        """Convert an EPUB e-book to PDF (via HTML → PDF pipeline)."""
        try:
            src = resolve_input(input_path)
            if not output_path:
                output_path = make_output_path(input_path, "", ".pdf")
            dst = resolve_output(output_path)

            html_content = _epub_to_html_str(src)
            try:
                from weasyprint import HTML
                HTML(string=html_content).write_pdf(str(dst))
                renderer = "weasyprint"
            except ImportError:
                from tools.pdf_import_tools import _html_to_pdf_reportlab
                _html_to_pdf_reportlab(html_content, str(dst))
                renderer = "reportlab"

            return {"success": True, "output_path": str(dst), "renderer": renderer}
        except ImportError:
            return {"success": False, "error": "ebooklib not installed. Run: pip install ebooklib"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ── RTF → TXT ──────────────────────────────────────────────────────────
    @app.tool()
    def rtf_to_txt(input_path: str, output_path: Optional[str] = None) -> Dict:
        """Extract plain text from an RTF (Rich Text Format) document."""
        try:
            from striprtf.striprtf import rtf_to_text
            src = resolve_input(input_path)
            if not output_path:
                output_path = make_output_path(input_path, "", ".txt")
            dst = resolve_output(output_path)
            rtf_content = src.read_text(encoding="utf-8", errors="replace")
            text = rtf_to_text(rtf_content)
            dst.write_text(text, encoding="utf-8")
            return {"success": True, "output_path": str(dst), "char_count": len(text)}
        except ImportError:
            return {"success": False, "error": "striprtf not installed. Run: pip install striprtf"}
        except Exception as e:
            return {"success": False, "error": str(e)}

    # ── RTF → HTML ─────────────────────────────────────────────────────────
    @app.tool()
    def rtf_to_html(input_path: str, output_path: Optional[str] = None) -> Dict:
        """Convert an RTF document to HTML (extracts text, basic structure)."""
        try:
            from striprtf.striprtf import rtf_to_text
            src = resolve_input(input_path)
            if not output_path:
                output_path = make_output_path(input_path, "", ".html")
            dst = resolve_output(output_path)
            rtf_content = src.read_text(encoding="utf-8", errors="replace")
            text = rtf_to_text(rtf_content)
            lines = "\n".join(
                f"<p>{html_module.escape(l)}</p>" if l.strip() else "<br>"
                for l in text.split("\n")
            )
            html_content = (
                f'<!DOCTYPE html>\n<html lang="en"><head><meta charset="UTF-8">'
                f'<title>{html_module.escape(src.stem)}</title>'
                f'<style>body{{font-family:sans-serif;max-width:800px;margin:auto;padding:2em}}</style>'
                f'</head><body>{lines}</body></html>'
            )
            dst.write_text(html_content, encoding="utf-8")
            return {"success": True, "output_path": str(dst), "char_count": len(text)}
        except ImportError:
            return {"success": False, "error": "striprtf not installed. Run: pip install striprtf"}
        except Exception as e:
            return {"success": False, "error": str(e)}
